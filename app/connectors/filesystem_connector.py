"""FileSystem knowledge connector for ingesting documents from local/shared drives."""

from __future__ import annotations

import hashlib
import os
from dataclasses import dataclass, field
from pathlib import Path

from app.connectors.base_connector import KnowledgeConnector
from app.document.document import Document


# File extensions supported by the connector
_SUPPORTED_EXTENSIONS = {".txt", ".md", ".html", ".pdf", ".docx", ".csv"}


@dataclass
class FileSystemConfig:
    """
    Configuration for the FileSystemConnector.

    Attributes:
        root_path: Root directory to scan
        extensions: File extensions to include (defaults to all supported)
        recursive: Whether to recurse into subdirectories
        max_file_size_bytes: Skip files larger than this limit (default 10 MB)
        exclude_patterns: Glob-style directory or file name patterns to skip
    """

    root_path: str
    extensions: set[str] = field(
        default_factory=lambda: set(_SUPPORTED_EXTENSIONS),
    )
    recursive: bool = True
    max_file_size_bytes: int = 10 * 1024 * 1024  # 10 MB
    exclude_patterns: list[str] = field(default_factory=list)


class FileSystemConnector(KnowledgeConnector):
    """
    Scans a directory tree and converts supported files into Document objects.

    Supports: .txt, .md, .html, .pdf, .docx, .csv

    Usage:
        config = FileSystemConfig(root_path="/data/knowledge-base")
        async with FileSystemConnector(config) as connector:
            documents = await connector.fetch()
    """

    CONNECTOR_ID = "filesystem"

    def __init__(self, config: FileSystemConfig):
        self._config = config
        self._root = Path(config.root_path)

    @property
    def connector_id(self) -> str:
        return self.CONNECTOR_ID

    async def connect(self) -> None:
        """Validate the root path exists and is accessible."""
        if not self._root.exists():
            raise FileNotFoundError(
                f"FileSystemConnector: root path does not exist: {self._root}"
            )
        if not self._root.is_dir():
            raise NotADirectoryError(
                f"FileSystemConnector: root path is not a directory: {self._root}"
            )

    async def disconnect(self) -> None:
        """No persistent resources to release."""
        pass

    async def fetch(self) -> list[Document]:
        """
        Walk the directory tree and load all supported files as Documents.

        Returns:
            List of Document objects, one per supported file
        """
        documents: list[Document] = []

        for file_path in self._discover_files():
            document = self._load_file(file_path)
            if document is not None:
                documents.append(document)

        return documents

    # -------------------------------------------------------------------------
    # Private helpers
    # -------------------------------------------------------------------------

    def _discover_files(self) -> list[Path]:
        """Return all matching file paths under root."""
        files: list[Path] = []

        if self._config.recursive:
            candidates = self._root.rglob("*")
        else:
            candidates = self._root.glob("*")

        for path in candidates:
            if not path.is_file():
                continue

            # Extension filter
            if path.suffix.lower() not in self._config.extensions:
                continue

            # Size filter
            try:
                if path.stat().st_size > self._config.max_file_size_bytes:
                    continue
            except OSError:
                continue

            # Exclude pattern filter
            if self._is_excluded(path):
                continue

            files.append(path)

        return sorted(files)

    def _is_excluded(self, path: Path) -> bool:
        """Return True if the path matches any exclude pattern."""
        for pattern in self._config.exclude_patterns:
            # Check against filename and all parent directory names
            if path.match(pattern):
                return True
            for parent in path.relative_to(self._root).parts:
                if Path(parent).match(pattern):
                    return True
        return False

    def _load_file(self, path: Path) -> Document | None:
        """Load a single file and return a Document, or None on failure."""
        suffix = path.suffix.lower()

        try:
            if suffix in {".txt", ".md", ".html", ".csv"}:
                text = self._read_text_file(path)
            elif suffix == ".pdf":
                text = self._read_pdf(path)
            elif suffix == ".docx":
                text = self._read_docx(path)
            else:
                return None
        except Exception:
            return None

        if not text or not text.strip():
            return None

        return Document(
            id=self._make_document_id(path),
            name=path.name,
            text=text.strip(),
            metadata={
                "source": "filesystem",
                "path": str(path),
                "extension": suffix,
                "size_bytes": str(path.stat().st_size),
            },
        )

    @staticmethod
    def _read_text_file(path: Path) -> str:
        """Read a plain-text file with UTF-8 (fallback to latin-1)."""
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return path.read_text(encoding="latin-1")

    @staticmethod
    def _read_pdf(path: Path) -> str:
        """Extract text from a PDF using pypdf."""
        try:
            import pypdf

            reader = pypdf.PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        except ImportError:
            raise RuntimeError(
                "pypdf is required for PDF support. "
                "Install it with: pip install pypdf"
            )

    @staticmethod
    def _read_docx(path: Path) -> str:
        """Extract text from a DOCX using python-docx."""
        try:
            import docx

            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text)
        except ImportError:
            raise RuntimeError(
                "python-docx is required for DOCX support. "
                "Install it with: pip install python-docx"
            )

    @staticmethod
    def _make_document_id(path: Path) -> str:
        """Create a stable document ID from the file path."""
        return hashlib.sha256(str(path).encode()).hexdigest()[:16]
