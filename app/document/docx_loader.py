from pathlib import Path

from docx import Document as DocxDocument

from app.document.document import Document
from app.document.loaders.document_loader import DocumentLoader


class DocxLoader(DocumentLoader):
    """
    Loads Microsoft Word (.docx) documents.
    """

    def load(
        self,
        path: str,
    ) -> Document:

        file = Path(path)

        doc = DocxDocument(path)

        paragraphs = []

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        document_text = "\n\n".join(paragraphs)

        return Document(
            id=file.stem,
            name=file.name,
            text=document_text,
            metadata={
                "document_type": "docx",
                "extension": ".docx",
                "paragraph_count": str(len(doc.paragraphs)),
                "source": file.name,
            },
        )